''' module for processing analysis files '''

import pandas as pd
import statsmodels.api as sm
from statsmodels.formula.api import ols
from statsmodels.stats.multicomp import pairwise_tukeyhsd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import seaborn as sns
import os

##################################################
# functions for renaming column names for analysis
##################################################

def rename_teams_trial_summary_profiles_surveys_for_analysis(teams_trial_summary_profiles_surveys_file_path,
                                                             output_file_path):
    df = pd.read_csv(teams_trial_summary_profiles_surveys_file_path)
    print(df.columns)

    out_df = df.copy()
    out_df.to_csv(output_file_path)

#########################################
# functions for individual analyses ANOVA
#########################################

# Define the color palette
black_gold_palette = ["#000000", "#FFD700"]  # Black and gold colors
sns.set_palette(sns.color_palette(black_gold_palette))
taskwork_category_order = ['Low Taskwork Pot', 'High Taskwork Pot']
teamwork_category_order = ['Low Teamwork Pot', 'High Teamwork Pot']

def perform_analysis_and_generate_document(dependent_var, df, individual_players_analysis_dir_path, combined_doc):
    df = df.copy()

    # Exclude rows with NaN values in the dependent variable or in the factors used for ANOVA
    df.dropna(subset=[dependent_var, 'Teamwork_potential_category', 'Taskwork_potential_category_conservative'],
              inplace=True)

    # Adjust DataFrame for plotting
    df['Teamwork potential category'] = df['Teamwork_potential_category'].str.replace('_', ' ')
    df['Taskwork potential category'] = df['Taskwork_potential_category_conservative'].str.replace('_', ' ')

    # Generating split violin plot with seaborn
    plt.figure(figsize=(12, 8))
    sns.violinplot(x='Teamwork potential category', y=dependent_var, hue='Taskwork potential category', data=df,
                   palette=black_gold_palette, split=True, inner=None,
               order=teamwork_category_order, hue_order=taskwork_category_order)
    plt.title(f'{dependent_var.replace("_", " ")} by Teamwork and Taskwork Potential Category (Split)', fontsize = 28)
    plt.xlabel('Teamwork Potential Category', fontsize = 20)
    plt.ylabel(dependent_var.replace('_', ' '), fontsize = 20)
    plt.xticks(rotation=45, fontsize=20)
    plt.yticks(fontsize=20)
    plt.tight_layout()
    plt.legend(title='Taskwork Potential Category', bbox_to_anchor=(1.05, 1), loc='upper left', fontsize = 20, title_fontsize=20)
    split_plot_file_name = f'{dependent_var}_split_violinplot.png'
    plt.savefig(split_plot_file_name, bbox_inches='tight')
    plt.close()

    # Generating full violin plot with seaborn
    plt.figure(figsize=(12, 8))
    sns.violinplot(x='Teamwork potential category', y=dependent_var, hue='Taskwork potential category', data=df,
                   palette=black_gold_palette, inner=None,
               order=teamwork_category_order, hue_order=taskwork_category_order)
    plt.title(f'{dependent_var.replace("_", " ")} by Teamwork and Taskwork Potential Category (Full)', fontsize = 28)
    plt.xlabel('Teamwork Potential Category', fontsize = 20)
    plt.ylabel(dependent_var.replace('_', ' '), fontsize = 20)
    plt.xticks(rotation=45, fontsize=20)
    plt.yticks(fontsize=20)
    plt.tight_layout()
    plt.legend(title='Taskwork Potential Category', bbox_to_anchor=(1.05, 1), loc='upper left', fontsize = 20, title_fontsize=20)
    full_plot_file_name = f'{dependent_var}_full_violinplot.png'
    plt.savefig(full_plot_file_name, bbox_inches='tight')
    plt.close()

    # Generating boxplot with seaborn for better visuals and colorblind-friendly palette
    plt.figure(figsize=(24, 12))
    boxplot = sns.boxplot(x='Teamwork potential category', y=dependent_var, hue='Taskwork potential category',
                          data=df, palette=black_gold_palette, order=teamwork_category_order, hue_order=taskwork_category_order,
                      showmeans=True, showfliers=False, showcaps=False,
                      meanprops={"marker": "o", "markerfacecolor": "white", "markeredgecolor": "black", "markersize": "10"},
                      medianprops={"color": "white", "linewidth": 2})
    plt.title(f'{dependent_var.replace("_", " ")} Boxplot by Teamwork and Taskwork Potential Category', fontsize=30, pad=40)
    plt.xlabel('Teamwork Potential Category', fontsize=24)
    plt.ylabel(dependent_var.replace('_', ' '), fontsize=24)
    plt.xticks(rotation=0, fontsize=20)
    plt.yticks(fontsize=20)
    plt.tight_layout()
    plt.legend(title='Taskwork Potential Category', bbox_to_anchor=(1.05, 1), loc='upper left', fontsize=20, title_fontsize=20)
    plt.subplots_adjust(top=2)  # Provide more space at the top
    # Use tight_layout with additional padding
    plt.tight_layout(pad=4.0)  # Increase padding to ensure elements are not overlapping
    # Saving the boxplot
    boxplot_file_name = f'{dependent_var}_boxplot.png'
    plt.savefig(boxplot_file_name, bbox_inches='tight')
    plt.close()

    # Perform two-way ANOVA
    formula = f'{dependent_var} ~ Teamwork_potential_category + Taskwork_potential_category_conservative + Teamwork_potential_category:Taskwork_potential_category_conservative'
    model = ols(formula, data=df).fit()
    anova_results = sm.stats.anova_lm(model, typ=2)

    # Calculate eta squared for effect size and add it to the ANOVA results
    anova_results['eta_sq'] = anova_results.sum_sq / sum(anova_results.sum_sq)

    # Simplify the names for the output
    anova_results.index = anova_results.index.str.replace("C\((.*?)\)", r"\1", regex=True)


    # Add analysis content to the combined document
    combined_doc.add_heading(f'Two-Way ANOVA Results with Post-Hoc Analyses for {dependent_var}', level=1)

    # Adding the ANOVA table to the combined document
    combined_doc.add_heading('ANOVA Table', level=2)
    table = combined_doc.add_table(rows=1, cols=len(anova_results.columns) + 1, style='Table Grid')
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Source'
    for i, col_name in enumerate(anova_results.columns):
        hdr_cells[i + 1].text = col_name

    for index, row in anova_results.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(index)
        for i, value in enumerate(row):
            row_cells[i + 1].text = f'{value:.4f}'


    # Create a new Word document
    doc = Document()
    doc.add_heading(f'Two-Way ANOVA Results with Post-Hoc Analyses for {dependent_var}', 0)

    # ANOVA table
    doc.add_heading('ANOVA Table', level=1)
    anova_table = doc.add_table(rows=anova_results.shape[0] + 1, cols=len(anova_results.columns) + 1)
    anova_table.style = 'Table Grid'

    # Header row
    anova_header_cells = anova_table.rows[0].cells
    anova_header_cells[0].text = 'Source'
    for i, col_name in enumerate(anova_results.columns):
        anova_header_cells[i + 1].text = col_name

    # Data rows
    for index, row in enumerate(anova_results.itertuples(), start=1):
        cells = anova_table.rows[index].cells
        cells[0].text = row.Index
        for i, value in enumerate(row[1:], start=1):
            cells[i].text = f'{value:.4f}'

    # Conduct and format post-hoc comparisons, including interaction groupings
    for var in ['Teamwork_potential_category', 'Taskwork_potential_category_conservative', 'Interaction']:
        mc = pairwise_tukeyhsd(df[dependent_var], df[var], alpha=0.05)
        result_summary = mc.summary()

        # Convert result summary to a DataFrame for easier handling
        result_df = pd.DataFrame(result_summary.data[1:], columns=result_summary.data[0])

        #for individual document
        doc.add_heading(f'Post-hoc for {var} (Bonferroni corrected)', level=2)
        post_hoc_table = doc.add_table(rows=len(result_df) + 1, cols=len(result_df.columns))
        post_hoc_table.style = 'Table Grid'

        # For combined document
        combined_doc.add_heading(f'Post-hoc for {var} (Bonferroni corrected)', level=2)
        combined_post_hoc_table = combined_doc.add_table(rows=len(result_df) + 1, cols=len(result_df.columns))
        combined_post_hoc_table.style = 'Table Grid'

        # Add header row
        for i, column in enumerate(result_df.columns):
            post_hoc_table.cell(0, i).text = column
            combined_post_hoc_table.cell(0, i).text = column

        # Fill table with data
        for i, row in enumerate(result_df.itertuples(), start=1):
            for j, value in enumerate(row[1:]):  # Skip the index
                post_hoc_table.cell(i, j).text = str(value)
                combined_post_hoc_table.cell(i, j).text = str(value)


    #doc.add_heading(f'Two-Way ANOVA Results with Post-Hoc Analyses for {dependent_var}', 0)
    # Add ANOVA table, post-hoc analyses, and the split violin plot to the document
    doc.add_page_break()
    doc.add_heading('Split Violin Plot Analysis', level=1)
    doc.add_picture(split_plot_file_name, width=Inches(6))


    # Add the full violin plot to the document
    doc.add_page_break()
    doc.add_heading('Full Violin Plot Analysis', level=1)
    doc.add_picture(full_plot_file_name, width=Inches(6))

    # Insert boxplot into the Word document
    doc.add_page_break()
    doc.add_heading('Boxplot Analysis', level=1)
    doc.add_picture(boxplot_file_name, width=Inches(6))

    # Save the document
    # output_path = f'C:\\Post-doc Work\\ASIST Study 4\\Analysis\\IndividualPlayers\\Study_4_individualProfiles_{dependent_var.replace(" ", "_")}_ANOVA_Results.docx'
    output_path = os.path.join(individual_players_analysis_dir_path, f'individual_profiles_{dependent_var.replace(" ", "_")}_ANOVA_results.docx')
    doc.save(output_path)
    print(f'Document saved to {output_path}')


    # Add the split violin plot to the combined document
    combined_doc.add_page_break()
    combined_doc.add_heading('Split Violin Plot Analysis', level=2)
    combined_doc.add_picture(split_plot_file_name, width=Inches(6))
    os.remove(split_plot_file_name)  # Clean up the saved plot image

    # Add the full violin plot to the combined document
    combined_doc.add_page_break()
    combined_doc.add_heading('Full Violin Plot Analysis', level=2)
    combined_doc.add_picture(full_plot_file_name, width=Inches(6))
    os.remove(full_plot_file_name)  # Clean up the saved plot image

    # Add the boxplot to the combined document
    combined_doc.add_page_break()
    combined_doc.add_heading('Boxplot Analysis', level=2)
    combined_doc.add_picture(boxplot_file_name, width=Inches(6))
    os.remove(boxplot_file_name)  # Clean up the saved plot image


def write_individual_analyses_anova(individual_player_profiles_team_alignment_trial_measures_for_analysis_file_path,
                                    individual_players_analysis_dir_path,
                                    combined_doc_file_path):
    # Initialize the combined document outside the loop
    # combined_doc_path = 'C:\\Post-doc Work\\ASIST Study 4\\Analysis\\IndividualPlayers\\Study_4_playerProfiles_ANOVA_Results_CombinedAnalyses.docx'
    combined_doc = Document()

    # Read the CSV file
    # csv_file_path = 'C:\\Post-doc Work\\ASIST Study 4\\Study_4_individual_playerProfiles_teamAlignment_trialMeasures_ForAnalysis.csv'
    df = pd.read_csv(individual_player_profiles_team_alignment_trial_measures_for_analysis_file_path)

    # Adding interaction term as a new column
    df['Interaction'] = df['Teamwork_potential_category'] + '_' + df['Taskwork_potential_category_conservative']

    # List of dependent variables to analyze
    #dependent_variables = ['TeammatesRescued','TimesFrozen','TextChatsSent',	'FlagsPlaced',	'FiresExtinguished',	'DamageTaken',	'NumCompletePostSurveys',	'BombBeaconsPlaced',	'BudgetExpended',	'CommBeaconsPlaced',	'BOMBS_EXPLODED_CHAINED',	'BOMBS_EXPLODED_FIRE',	'BOMBS_EXPLODED_STANDARD',	'BOMBS_DEFUSED_CHAINED',	'BOMBS_DEFUSED_FIRE',	'BOMBS_DEFUSED_STANDARD', 'SATIS_score',	'SATIS_workedTogether',	'SATIS_teamPlan',	'SATIS_teamAgain',	'SATIS_teamCapable',	'EFF_workEthic',	'EFF_overcomeProblems',	'EFF_planStrategy',	'EFF_maintainPositivity',	'EFF_disposeBombs',	'EFF_speedRun',	'EFF_knowledgeCoord',	'EFF_roleCoord',	'advisorEVAL_improvedScore',	'AdvisorEVAL_improvedCoord',	'advisorEVAL_comfortDependingOn',	'advisorEVAL_understoodRecommends',	'advisorEVAL_wasTrustworthy',	'PRSS_transition',	'PRSS_action',	'PRSS_interpersonal',	'PRSS_avg']   # Extend this list with other dependent variables

    dependent_variables = ['Text_Chats_Sent',
                        'Bomb_Beacons_Placed',
                        'Bombs_Defused_Chained',
                        'Bombs_Defused_Standard',
                        'Bombs_Exploded_Chained',
                        'Bombs_Exploded_Standard',
                        'Team_Satisfaction_Rating_Team_Score',
                        'Team_Satisfaction_Rating_Team_Plan',
                        'Team_Efficacy_Rating_Team_Strategy',
                        'Team_Efficacy_Rating_Team_Knowledge_Coordination',
                        'Team_Efficacy_Rating_Team_Role_Coordination',
                        'ASI_Evaluation_Improved_Score',
                        'ASI_Evaluation_Improved_Coordination',
                        'ASI_Evaluation_Dependability',
                        'ASI_Evaluation_Was_Trustworthy',
                        'Team_Process_Rating_average'
                        ]


    for dependent_var in dependent_variables:
        perform_analysis_and_generate_document(dependent_var, df, individual_players_analysis_dir_path, combined_doc)

    # After generating all individual documents, combine them into a single document
    # output_directory = 'C:\\Post-doc Work\\ASIST Study 4\\Analysis\\IndividualPlayers'

    # Save the combined document after all analyses
    combined_doc.save(combined_doc_file_path)